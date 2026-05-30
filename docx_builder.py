"""
docx_builder.py — Qarta Legal document formatter v3.

Produces three lawyer-ready Word documents per adaptation job:
  - clean.docx      : final adapted contract
  - redline.docx    : change-tracked diff (strikethrough red / underline green)
  - commentary.docx : clause-by-clause legal analysis

Typography:
  Body            — Times New Roman 10pt  #1A1A2E
  Labels/badges   — Arial
  Metadata tags   — Courier New
  Clause titles   — Times New Roman bold small-caps
  Clause numbers  — Times New Roman bold amethyst #7C3AED
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ══════════════════════════════════════════════════════════════════════════════
# Lookup tables
# ══════════════════════════════════════════════════════════════════════════════

DOC_TYPE_LABELS = {
    "nda":                                   "Non-Disclosure Agreement",
    "non_disclosure_agreement":              "Non-Disclosure Agreement",
    "jv_agreement":                          "Joint Venture Agreement",
    "joint_venture":                         "Joint Venture Agreement",
    "jv":                                    "Joint Venture Agreement",
    "mou":                                   "Memorandum of Understanding",
    "memorandum_of_understanding":           "Memorandum of Understanding",
    "service_agreement":                     "Service Agreement",
    "employment_contract":                   "Employment Contract",
    "employment":                            "Employment Contract",
    "exclusive_distribution":                "Exclusive Distribution Agreement",
    "exclusive_distribution_agreement":      "Exclusive Distribution Agreement",
    "non_exclusive_distribution":            "Non-Exclusive Distribution Agreement",
    "non_exclusive_distribution_agreement":  "Non-Exclusive Distribution Agreement",
    "sha":                                   "Shareholders Agreement",
    "shareholders_agreement":                "Shareholders Agreement",
    "ip_licence":                            "IP Licence Agreement",
    "ip_license":                            "IP Licence Agreement",
    "franchise":                             "Franchise Agreement",
    "franchise_agreement":                   "Franchise Agreement",
    "pdpa":                                  "PDPA Data Protection Agreement",
}

CORRIDOR_LABELS = {
    "CN_SG": "China → Singapore",
    "SG_ID": "Singapore → Indonesia",
    "CN_ID": "China → Indonesia",
    "SG_MY": "Singapore → Malaysia",
    "CN_MY": "China → Malaysia",
}

CORRIDOR_GOV_LAW = {
    "CN_SG": "Singapore Law",
    "SG_ID": "Indonesian Law",
    "CN_ID": "Indonesian Law",
    "SG_MY": "Malaysian Law",
    "CN_MY": "Malaysian Law",
}

# ══════════════════════════════════════════════════════════════════════════════
# Colour palette
# ══════════════════════════════════════════════════════════════════════════════

AMETHYST   = RGBColor(0x7C, 0x3A, 0xED)   # #7C3AED — brand accent
BODY_CLR   = RGBColor(0x1A, 0x1A, 0x2E)   # #1A1A2E — primary body
SECONDARY  = RGBColor(0x6B, 0x72, 0x80)   # #6B7280 — grey labels
BLUE_GREY  = RGBColor(0x37, 0x51, 0x6F)   # #37516F — reason text
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
# Flag colours
ACTION_CLR = RGBColor(0xF5, 0x9E, 0x0B)   # #F59E0B amber border
LAWYER_CLR = RGBColor(0x93, 0x33, 0xEA)   # #9333EA purple border
NOTE_CLR   = RGBColor(0x0E, 0xA5, 0xE9)   # #0EA5E9 blue border
# Diff colours
INS_CLR    = RGBColor(0x05, 0x96, 0x69)   # #059669 green — insertions
DEL_CLR    = RGBColor(0xDC, 0x26, 0x26)   # #DC2626 red  — deletions

# Hex strings (no #) for XML attributes
AMETHYST_HX   = "7C3AED"
ACTION_BG_HX  = "FFF3E0"
ACTION_BD_HX  = "F59E0B"
LAWYER_BG_HX  = "FDF2F8"
LAWYER_BD_HX  = "9333EA"
NOTE_BG_HX    = "F0F9FF"
NOTE_BD_HX    = "0EA5E9"
COVER_BG_HX   = "F5F3FF"   # near-white amethyst tint
TABLE_HDR_HX  = "F3F0FF"   # light purple table header
GREY_BD_HX    = "E5E7EB"   # light grey for header/footer borders
DIVIDER_HX    = "DDD6FE"   # muted amethyst for clause dividers
HIGH_BG_HX    = "FEEAEA"   # light red — HIGH severity rows
LOW_BG_HX     = "F3F4F6"   # light grey — LOW severity rows

# ══════════════════════════════════════════════════════════════════════════════
# Typography
# ══════════════════════════════════════════════════════════════════════════════

FONT_BODY = "Times New Roman"
FONT_UI   = "Arial"
FONT_MONO = "Courier New"

SZ_BODY   = Pt(10)
SZ_SMALL  = Pt(8)
SZ_LABEL  = Pt(9)
SZ_TITLE  = Pt(20)
SZ_BRAND  = Pt(9)
SZ_H1     = Pt(11)
SZ_H2     = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# Regex
# ══════════════════════════════════════════════════════════════════════════════

RE_PLACEHOLDER        = re.compile(r'(\[[A-Z][A-Z0-9 _]+:[^\]]{1,120}\])')
RE_TOP_CLAUSE         = re.compile(r'^(\d+)\.\s{1,4}([A-Z].*)$')
RE_SUMMARY_TABLE_START = re.compile(r'^=== SUMMARY TABLE ===', re.IGNORECASE)
RE_SUMMARY_TABLE_END   = re.compile(r'^=== END SUMMARY TABLE ===', re.IGNORECASE)
RE_TABLE_ROW           = re.compile(r'^\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|$')
RE_SUB_CLAUSE    = re.compile(r'^(\d+\.\d+)\s{1,4}(\S.*)$')
RE_SUB_SUB       = re.compile(r'^(\d+\.\d+\.\d+)\s{1,4}(\S.*)$')
RE_SEPARATOR     = re.compile(r'^[━─=\-]{5,}$')
RE_CLAUSE_HEADER = re.compile(r'^\[CLAUSE\s+[\d.]+', re.IGNORECASE)
RE_PDPA_ITEM     = re.compile(r'^[✅⚠️☑]')
RE_ACTION_FLAG   = re.compile(r'^\s*[⚠*]\s*ACTION REQUIRED', re.IGNORECASE)
RE_LAWYER_FLAG   = re.compile(r'^\s*[⚠*┌│]\s*LAWYER REVIEW', re.IGNORECASE)
RE_NOTE_FLAG     = re.compile(r'^\s*[ℹ*]\s*NOTE[:\s]', re.IGNORECASE)
RE_BOX_START     = re.compile(r'^[┌╔]')
RE_BOX_END       = re.compile(r'^[└╚]')
RE_BOX_MID       = re.compile(r'^[│║]')

# Markdown inline — matches **bold** and *italic* spans (bold checked first
# so ** is never mis-split as two * italic markers)
_RE_INLINE_MD  = re.compile(r'\*\*(.+?)\*\*|\*([^*\n]+?)\*', re.DOTALL)

# Markdown structural
RE_MD_HEADING  = re.compile(r'^(#{1,3})\s+(.+)$')   # # / ## / ###
RE_MD_HR       = re.compile(r'^-{3}$')               # --- only (not ────)
RE_MD_BLOCK    = re.compile(r'^>\s*(.*)')             # > blockquote

FIELD_LABELS = ('Original:', 'Change:', 'Reason:', 'Action required:')

# Bilingual (SG-ID) — corridors that get a two-column EN | local-language layout
_BILINGUAL_CORRIDORS  = frozenset({"SG_ID"})
_SWORN_PLACEHOLDER    = "[ Sworn translation required — penerjemah tersumpah ]"
_ID_MARKERS           = (
    'KUHPerdata', 'Pasal ', 'UU Bahasa', 'UU No.', 'Bea Materai',
    'penerjemah', 'tersumpah', 'perjanjian',
)
_RE_ID_BLOCK_START = re.compile(r'^\[BAHASA INDONESIA\]',  re.IGNORECASE)
_RE_ID_BLOCK_END   = re.compile(r'^\[/BAHASA INDONESIA\]', re.IGNORECASE)


# ══════════════════════════════════════════════════════════════════════════════
# Low-level XML / run helpers
# ══════════════════════════════════════════════════════════════════════════════

def _fmt(run, font=FONT_BODY, bold=False, italic=False, strike=False,
         underline=False, small_caps=False, color=None, size=None):
    run.font.name       = font
    run.font.bold       = bold
    run.font.italic     = italic
    run.font.strike     = strike
    run.font.underline  = underline
    run.font.small_caps = small_caps
    run.font.size       = size or SZ_BODY
    run.font.color.rgb  = color or BODY_CLR


def _add_md_runs(para, text: str, font=FONT_BODY, color=None, italic=False,
                 bold=False, strike=False, underline=False, size=None):
    """Parse **bold** and *italic* markdown within `text` and emit styled runs.

    Base `italic` / `bold` kwargs are additive: a **bold** span inside an
    already-italic field (e.g. Original:) becomes bold-italic.  Unrecognised
    asterisks that don't form a valid span are passed through verbatim.
    """
    last = 0
    for m in _RE_INLINE_MD.finditer(text):
        # Normal text before this match
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            _fmt(r, font=font, bold=bold, italic=italic, strike=strike,
                 underline=underline, color=color, size=size)
        if m.group(1) is not None:       # **bold**
            r = para.add_run(m.group(1))
            _fmt(r, font=font, bold=True, italic=italic, strike=strike,
                 underline=underline, color=color, size=size)
        else:                            # *italic*
            r = para.add_run(m.group(2))
            _fmt(r, font=font, bold=bold, italic=True, strike=strike,
                 underline=underline, color=color, size=size)
        last = m.end()
    # Remaining text after the last match (or the whole string if no matches)
    if last < len(text):
        r = para.add_run(text[last:])
        _fmt(r, font=font, bold=bold, italic=italic, strike=strike,
             underline=underline, color=color, size=size)


def _para_spacing(para, before=0, after=80):
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    para._p.get_or_add_pPr().append(sp)


def _set_indent(para, left_twips=0, hanging_twips=0):
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(left_twips))
    if hanging_twips:
        ind.set(qn('w:hanging'), str(hanging_twips))
    para._p.get_or_add_pPr().append(ind)


def _set_para_bg(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)


def _set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _set_cell_no_border(cell):
    """Remove all borders from a table cell (for layout tables)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _add_para_border(para, sides, color_hex, thickness='4', space='4'):
    """Add borders to specified sides of a paragraph."""
    pPr = para._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    thickness)
        el.set(qn('w:space'), space)
        el.set(qn('w:color'), color_hex)
        bdr.append(el)
    pPr.append(bdr)


def _add_left_border(para, bdr_hex, bg_hex=None, thickness='24'):
    """Left-accent flag box — thick left border, optional background fill."""
    _add_para_border(para, ['left'], bdr_hex, thickness=thickness, space='8')
    if bg_hex:
        _set_para_bg(para, bg_hex)


def _run_badge(para, text, bg_hex=AMETHYST_HX, fg_color=None):
    """Inline colored-background badge run."""
    run = para.add_run(f"  {text}  ")
    run.font.name  = FONT_UI
    run.font.size  = SZ_SMALL
    run.font.bold  = True
    run.font.color.rgb = fg_color or WHITE
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg_hex)
    rPr.append(shd)
    return run


def _add_page_number(para):
    run = para.add_run()
    for tag, val in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = val
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)


def _new_doc() -> Document:
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin    = Inches(0.9)
    sect.bottom_margin = Inches(0.9)
    sect.left_margin   = Inches(1.2)
    sect.right_margin  = Inches(1.2)
    doc.styles['Normal'].font.name = FONT_BODY
    doc.styles['Normal'].font.size = SZ_BODY
    return doc, sect


def _horizontal_rule(doc, color_hex=DIVIDER_HX):
    p = doc.add_paragraph()
    _para_spacing(p, before=60, after=60)
    _add_para_border(p, ['bottom'], color_hex, thickness='4', space='1')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# Page header / footer
# ══════════════════════════════════════════════════════════════════════════════

def _setup_header_footer(section, doc_type_label: str, corridor_label: str):
    """
    Header  left : "QARTA LEGAL · {doc_type} · {corridor}"   grey
    Header  right: "CONFIDENTIAL — FOR LAWYER REVIEW"          amethyst bold
            + thin grey bottom border

    Footer  left : disclaimer italic grey
    Footer  right: page number
            + thin grey top border
    """
    # ── PAGE HEADER ────────────────────────────────────────────────────────────
    hp = section.header.paragraphs[0]
    hp.clear()

    # Grey bottom border on header paragraph
    _add_para_border(hp, ['bottom'], GREY_BD_HX, thickness='4', space='4')

    # Right-edge tab stop (6 inches = 8640 twips)
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tr = OxmlElement('w:tab')
    tr.set(qn('w:val'), 'right')
    tr.set(qn('w:pos'), '8640')
    tabs.append(tr)
    pPr.append(tabs)

    rl = hp.add_run(f"QARTA LEGAL  ·  {doc_type_label}  ·  {corridor_label}")
    rl.font.name  = FONT_UI
    rl.font.size  = SZ_SMALL
    rl.font.color.rgb = SECONDARY

    hp.add_run('\t')

    rr = hp.add_run("CONFIDENTIAL — FOR LAWYER REVIEW")
    rr.font.name  = FONT_UI
    rr.font.size  = SZ_SMALL
    rr.font.bold  = True
    rr.font.color.rgb = AMETHYST

    # ── PAGE FOOTER ────────────────────────────────────────────────────────────
    fp = section.footer.paragraphs[0]
    fp.clear()

    # Grey top border on footer paragraph
    _add_para_border(fp, ['top'], GREY_BD_HX, thickness='4', space='4')

    pPr2 = fp._p.get_or_add_pPr()
    tabs2 = OxmlElement('w:tabs')
    tr2 = OxmlElement('w:tab')
    tr2.set(qn('w:val'), 'right')
    tr2.set(qn('w:pos'), '8640')
    tabs2.append(tr2)
    pPr2.append(tabs2)

    rd = fp.add_run(
        "AI-assisted adaptation. Not legal advice. Requires attorney review before execution."
    )
    rd.font.name   = FONT_UI
    rd.font.size   = SZ_SMALL
    rd.font.italic = True
    rd.font.color.rgb = SECONDARY

    fp.add_run('\t')
    _add_page_number(fp)


# ══════════════════════════════════════════════════════════════════════════════
# Cover block
# ══════════════════════════════════════════════════════════════════════════════

def _add_cover_block(doc, doc_type_label: str, corridor_label: str,
                     gov_law: str, variant: str, date_str: str):
    """
    Two-column layout table (no borders):
      LEFT  col: QARTA LEGAL (small caps amethyst) / title (Times New Roman 20pt bold) /
                 corridor + governing law (Arial 9pt grey)
      RIGHT col: AI-ASSISTED DRAFT badge / "For lawyer review..." italic /
                 preparation date Courier
    Separated from body by a 2pt amethyst bottom border.
    """
    tbl = doc.add_table(rows=1, cols=2)
    # Remove all table borders via tblPr
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tbl_pr)
    tbl_bdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tbl_bdr.append(el)
    tbl_pr.append(tbl_bdr)

    tbl.columns[0].width = Inches(3.6)
    tbl.columns[1].width = Inches(2.7)

    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    _set_cell_no_border(left_cell)
    _set_cell_no_border(right_cell)
    _set_cell_bg(left_cell,  COVER_BG_HX)
    _set_cell_bg(right_cell, COVER_BG_HX)

    # ── LEFT CELL ──────────────────────────────────────────────────────────────
    # "QARTA LEGAL" small-caps amethyst
    p_brand = left_cell.paragraphs[0]
    p_brand.clear()
    _para_spacing(p_brand, before=80, after=40)
    r_brand = p_brand.add_run("QARTA LEGAL")
    r_brand.font.name       = FONT_UI
    r_brand.font.size       = SZ_BRAND
    r_brand.font.bold       = True
    r_brand.font.small_caps = True
    r_brand.font.color.rgb  = AMETHYST

    # Document title Times New Roman 20pt bold dark
    p_title = left_cell.add_paragraph()
    _para_spacing(p_title, before=20, after=40)
    r_title = p_title.add_run(doc_type_label)
    r_title.font.name  = FONT_BODY
    r_title.font.size  = SZ_TITLE
    r_title.font.bold  = True
    r_title.font.color.rgb = BODY_CLR

    # Corridor + governing law subtitle — Arial 9pt grey
    p_sub = left_cell.add_paragraph()
    _para_spacing(p_sub, before=0, after=80)
    r_sub = p_sub.add_run(f"{corridor_label}  ·  Governed by {gov_law}")
    r_sub.font.name  = FONT_UI
    r_sub.font.size  = SZ_LABEL
    r_sub.font.color.rgb = SECONDARY

    # ── RIGHT CELL ─────────────────────────────────────────────────────────────
    right_cell.paragraphs[0].clear()

    # Spacer paragraph to push content down visually
    p_space = right_cell.paragraphs[0]
    _para_spacing(p_space, before=80, after=40)
    p_space.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # AI-ASSISTED DRAFT badge
    _run_badge(p_space, "AI-ASSISTED DRAFT")

    # "For lawyer review before execution" italic
    p_disc = right_cell.add_paragraph()
    _para_spacing(p_disc, before=20, after=20)
    p_disc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_disc = p_disc.add_run("For lawyer review before execution")
    r_disc.font.name   = FONT_UI
    r_disc.font.size   = SZ_SMALL
    r_disc.font.italic = True
    r_disc.font.color.rgb = SECONDARY

    # Preparation date Courier
    p_date = right_cell.add_paragraph()
    _para_spacing(p_date, before=0, after=80)
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_date = p_date.add_run(f"Prepared: {date_str}")
    r_date.font.name  = FONT_MONO
    r_date.font.size  = SZ_SMALL
    r_date.font.color.rgb = SECONDARY

    # Separator paragraph — 2pt amethyst bottom border
    p_sep = doc.add_paragraph()
    _para_spacing(p_sep, before=0, after=160)
    _add_para_border(p_sep, ['bottom'], AMETHYST_HX, thickness='16', space='4')
    _set_para_bg(p_sep, COVER_BG_HX)


# ══════════════════════════════════════════════════════════════════════════════
# Deal profile table
# ══════════════════════════════════════════════════════════════════════════════

def _add_deal_profile_table(doc, corridor: str, doc_type: str,
                             company_name: str, date_str: str):
    """
    Two-column table: label (bold grey Arial) / value (Times New Roman).
    Header row with light purple background.
    Employer entity value shown amber italic = user input required.
    """
    rows_data = [
        ("Corridor",         CORRIDOR_LABELS.get(corridor.upper(), corridor), 'normal'),
        ("Document type",    DOC_TYPE_LABELS.get(doc_type.lower(),
                             doc_type.replace('_', ' ').title()),            'normal'),
        ("Employer entity",  company_name or "[PARTY NAME — USER INPUT]",    'amber'),
        ("Governing law",    CORRIDOR_GOV_LAW.get(corridor.upper(), "—"),    'normal'),
        ("Adaptation date",  date_str,                                        'mono'),
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # approximate widths
    tbl.columns[0].width = Inches(1.8)
    tbl.columns[1].width = Inches(4.5)

    # Header row
    hdr = tbl.rows[0].cells
    _set_cell_bg(hdr[0], TABLE_HDR_HX)
    _set_cell_bg(hdr[1], TABLE_HDR_HX)
    for cell, txt in zip(hdr, ["Field", "Value"]):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(txt)
        r.font.name  = FONT_UI
        r.font.size  = SZ_LABEL
        r.font.bold  = True
        r.font.color.rgb = AMETHYST

    for field, value, style in rows_data:
        row = tbl.add_row().cells
        # Label cell
        pf = row[0].paragraphs[0]
        pf.clear()
        rf = pf.add_run(field)
        rf.font.name  = FONT_UI
        rf.font.size  = SZ_LABEL
        rf.font.bold  = True
        rf.font.color.rgb = SECONDARY
        # Value cell
        pv = row[1].paragraphs[0]
        pv.clear()
        rv = pv.add_run(value)
        if style == 'amber':
            rv.font.name   = FONT_BODY
            rv.font.size   = SZ_BODY
            rv.font.italic = True
            rv.font.color.rgb = ACTION_CLR
        elif style == 'mono':
            rv.font.name  = FONT_MONO
            rv.font.size  = SZ_BODY
            rv.font.color.rgb = BODY_CLR
        else:
            rv.font.name  = FONT_BODY
            rv.font.size  = SZ_BODY
            rv.font.color.rgb = BODY_CLR

    doc.add_paragraph()   # breathing room after table


# ══════════════════════════════════════════════════════════════════════════════
# Execution block
# ══════════════════════════════════════════════════════════════════════════════

def _add_execution_block(doc):
    _horizontal_rule(doc)

    p_h = doc.add_paragraph()
    _para_spacing(p_h, before=120, after=80)
    r_h = p_h.add_run("EXECUTION")
    r_h.font.name       = FONT_UI
    r_h.font.size       = SZ_H1
    r_h.font.bold       = True
    r_h.font.small_caps = True
    r_h.font.color.rgb  = AMETHYST

    parties = ["Employer / Party A", "Employee / Party B"]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'

    for i, party in enumerate(parties):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, TABLE_HDR_HX)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(party)
        r.font.name  = FONT_UI
        r.font.size  = SZ_LABEL
        r.font.bold  = True
        r.font.color.rgb = AMETHYST

    for label in ("Signature", "Name", "Title", "Date"):
        row = tbl.add_row().cells
        for cell in row:
            p = cell.paragraphs[0]
            p.clear()
            _para_spacing(p, before=120, after=20)
            rl = p.add_run(f"{label}:  ")
            rl.font.name  = FONT_UI
            rl.font.size  = SZ_LABEL
            rl.font.bold  = True
            rl.font.color.rgb = SECONDARY
            rb = p.add_run("_" * 36)
            rb.font.name  = FONT_BODY
            rb.font.size  = SZ_BODY
            rb.font.color.rgb = BODY_CLR


# ══════════════════════════════════════════════════════════════════════════════
# Clause heading helpers
# ══════════════════════════════════════════════════════════════════════════════

def _clause_parts(line: str):
    """Return (level, num_str, title_str) or (0,'','')."""
    m = RE_SUB_SUB.match(line)
    if m: return 3, m.group(1), m.group(2)
    m = RE_SUB_CLAUSE.match(line)
    if m: return 2, m.group(1), m.group(2)
    m = RE_TOP_CLAUSE.match(line)
    if m: return 1, m.group(1) + '.', m.group(2)
    return 0, '', ''


def _add_clause_heading(doc, num_str: str, title_str: str, level: int):
    """
    Amethyst Times New Roman bold clause number + dark small-caps Times New Roman title.
    Level 1: spacing + thin amethyst divider after.
    Level 2/3: indented with hanging indent.
    """
    para = doc.add_paragraph()
    sp_before = {1: 200, 2: 100, 3: 60}
    _para_spacing(para, before=sp_before.get(level, 100), after=60)

    # Hanging indent for sub-clauses: left=360 twips, hanging=360 twips
    if level == 2:
        _set_indent(para, left_twips=360, hanging_twips=360)
    elif level == 3:
        _set_indent(para, left_twips=720, hanging_twips=360)

    # Clause number — amethyst Times New Roman bold
    r_num = para.add_run(num_str + "  ")
    r_num.font.name  = FONT_BODY
    r_num.font.bold  = True
    r_num.font.size  = SZ_H1 if level == 1 else SZ_H2
    r_num.font.color.rgb = AMETHYST

    # Clause title — dark Times New Roman bold small-caps
    r_txt = para.add_run(title_str)
    r_txt.font.name       = FONT_BODY
    r_txt.font.bold       = (level <= 2)
    r_txt.font.small_caps = (level == 1)
    r_txt.font.size       = SZ_H1 if level == 1 else SZ_H2
    r_txt.font.color.rgb  = BODY_CLR

    if level == 1:
        _horizontal_rule(doc)

    return para


def _write_md_heading(doc, level: int, text: str):
    """Render a markdown # / ## / ### heading as a styled Word paragraph.

    Level 1 — large bold small-caps body font, amethyst divider after (mirrors
               top-level clause heading visual weight).
    Level 2 — bold body font, H2 size.
    Level 3 — bold body font, body size (used for sub-section titles).
    """
    para = doc.add_paragraph()
    if level == 1:
        _para_spacing(para, before=240, after=100)
        r = para.add_run(text)
        r.font.name       = FONT_BODY
        r.font.size       = SZ_H1
        r.font.bold       = True
        r.font.small_caps = True
        r.font.color.rgb  = BODY_CLR
        _horizontal_rule(doc)
    elif level == 2:
        _para_spacing(para, before=160, after=80)
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = SZ_H2
        r.font.bold  = True
        r.font.color.rgb = BODY_CLR
    else:   # level 3
        _para_spacing(para, before=100, after=60)
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = SZ_BODY
        r.font.bold  = True
        r.font.color.rgb = BODY_CLR
    return para


# ══════════════════════════════════════════════════════════════════════════════
# Flag box helpers
# ══════════════════════════════════════════════════════════════════════════════

def _detect_flag(line: str):
    if RE_ACTION_FLAG.match(line): return 'action'
    if RE_LAWYER_FLAG.match(line): return 'lawyer'
    if RE_NOTE_FLAG.match(line):   return 'note'
    return None


_FLAG_STYLES = {
    'action': (ACTION_BD_HX, ACTION_BG_HX, ACTION_CLR),
    'lawyer': (LAWYER_BD_HX, LAWYER_BG_HX, LAWYER_CLR),
    'note':   (NOTE_BD_HX,   NOTE_BG_HX,   NOTE_CLR),
}


def _add_flag_para(doc, text: str, flag_type: str, is_label=False):
    bdr_hex, bg_hex, clr = _FLAG_STYLES.get(flag_type, (AMETHYST_HX, COVER_BG_HX, AMETHYST))
    p = doc.add_paragraph()
    _para_spacing(p, before=20 if is_label else 0, after=20)
    _set_indent(p, left_twips=180)
    _add_left_border(p, bdr_hex, bg_hex=bg_hex)
    clean = text.lstrip('┌└│║⚠ℹ* ').strip()
    r = p.add_run(clean)
    r.font.name  = FONT_UI
    r.font.bold  = is_label
    r.font.size  = SZ_LABEL
    r.font.color.rgb = clr
    return p


# ══════════════════════════════════════════════════════════════════════════════
# CLEAN document
# ══════════════════════════════════════════════════════════════════════════════

def _para_with_placeholders(doc, line: str):
    """Body paragraph — [PLACEHOLDER:...] rendered bold amber Courier; **text** rendered bold."""
    para = doc.add_paragraph()
    _para_spacing(para, after=80)
    for part in RE_PLACEHOLDER.split(line):
        if not part:
            continue
        if RE_PLACEHOLDER.fullmatch(part):
            r = para.add_run(part)
            r.font.name  = FONT_MONO
            r.font.size  = SZ_BODY
            r.font.bold  = True
            r.font.color.rgb = ACTION_CLR
        else:
            _add_md_runs(para, part)
    return para


# ══════════════════════════════════════════════════════════════════════════════
# SG-ID bilingual layout helpers
# ══════════════════════════════════════════════════════════════════════════════

def _parse_bilingual_blocks(text: str) -> list:
    """Split clean text into per-clause bilingual blocks.

    Returns [{'en': str, 'id': str|None}, ...] where each dict covers one
    top-level clause.  Content between [BAHASA INDONESIA] / [/BAHASA INDONESIA]
    tags is routed to the 'id' slot; everything else goes to 'en'.
    If no [BAHASA INDONESIA] tags appear, 'id' is None for every block and the
    caller substitutes the sworn-translator placeholder.
    """
    blocks: list = []
    current_en: list = []
    current_id: list = []
    in_id       = False
    has_content = False

    def _flush():
        en  = '\n'.join(current_en).strip()
        id_ = '\n'.join(current_id).strip() or None
        if en or id_:
            blocks.append({'en': en, 'id': id_})

    for line in text.splitlines():
        s = line.strip()

        if _RE_ID_BLOCK_START.match(s):
            in_id = True
            continue
        if _RE_ID_BLOCK_END.match(s):
            in_id = False
            continue

        if in_id:
            current_id.append(line)
            has_content = True
            continue

        # Top-level clause boundary → flush previous block first
        if RE_TOP_CLAUSE.match(s) and has_content:
            _flush()
            current_en = [line]
            current_id = []
        else:
            current_en.append(line)
        has_content = True

    _flush()
    return blocks


def _render_lines_to_cell(container, lines_text: str) -> None:
    """Render clean-format lines into a table cell (or Document) container.

    Mirrors the _build_clean loop body so it works identically whether the
    target is a full Document or a table Cell — both expose .add_paragraph().
    The cell's initial auto-created empty paragraph acts as a small top-spacer.
    """
    flag_type = None
    for line in lines_text.splitlines():
        s = line.strip()
        if not s:
            flag_type = None
            continue
        if RE_SEPARATOR.match(s):
            continue
        if RE_BOX_END.match(s):
            flag_type = None
            continue

        # Markdown structural
        m_h = RE_MD_HEADING.match(s)
        if m_h:
            flag_type = None
            _write_md_heading(container, len(m_h.group(1)), m_h.group(2).strip())
            continue
        if RE_MD_HR.match(s):
            _horizontal_rule(container)
            continue
        m_bq = RE_MD_BLOCK.match(s)
        if m_bq:
            flag_type = None
            content = m_bq.group(1).strip()
            _add_flag_para(container, content or s, _detect_flag(content) or 'note')
            continue

        # Box-drawing flag blocks
        if RE_BOX_START.match(s):
            flag_type = _detect_flag(s) or 'lawyer'
            _add_flag_para(container, s, flag_type, is_label=True)
            continue
        if RE_BOX_MID.match(s) and flag_type:
            _add_flag_para(container, s, flag_type)
            continue

        # Inline flags
        detected = _detect_flag(s)
        if detected:
            flag_type = detected
            _add_flag_para(container, s, flag_type, is_label=True)
            continue
        if flag_type and not RE_TOP_CLAUSE.match(s):
            _add_flag_para(container, s, flag_type)
            continue
        flag_type = None

        # Clause headings
        level, num_str, title_str = _clause_parts(s)
        if level:
            _add_clause_heading(container, num_str, title_str, level)
            continue

        _para_with_placeholders(container, s)


def _add_bilingual_notice(doc) -> None:
    """Render the UU Bahasa Art. 31 compliance notice as an amber flag box."""
    p = doc.add_paragraph()
    _para_spacing(p, before=120, after=80)
    _set_indent(p, left_twips=180)
    _add_left_border(p, ACTION_BD_HX, bg_hex=ACTION_BG_HX, thickness='24')
    r = p.add_run(
        "⚠ BILINGUAL REQUIREMENT — UU Bahasa (Law No.…24/2009) "
        "Art.…31: Agreements involving Indonesian parties or performed in "
        "Indonesia must be executed in Bahasa Indonesia.  Where a bilingual version "
        "is used, the Bahasa Indonesia text is the legally prevailing version.  "
        "Clauses showing “" + _SWORN_PLACEHOLDER + "” require a sworn "
        "translator (penerjemah tersumpah) before execution."
    )
    r.font.name      = FONT_UI
    r.font.size      = SZ_LABEL
    r.font.color.rgb = ACTION_CLR


def _build_clean_bilingual(doc, text: str) -> None:
    """Render the clean document body as a two-column EN | BI bilingual table.

    LEFT  column (2.9 in) : English adapted clause text
    RIGHT column (2.9 in) : Bahasa Indonesia text produced by the engine, or
                             the sworn-translator placeholder in grey italic.

    Each top-level clause occupies one table row — keeping the two languages
    clause-aligned when the document is reviewed side-by-side.
    """
    _add_bilingual_notice(doc)

    blocks = _parse_bilingual_blocks(text)

    if not blocks:
        # Fallback: plain single-column rendering when text has no clause structure
        for line in text.splitlines():
            s = line.strip()
            if not s:
                continue
            lv, num_str, title_str = _clause_parts(s)
            if lv:
                _add_clause_heading(doc, num_str, title_str, lv)
            else:
                _para_with_placeholders(doc, s)
        return

    # Build table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(2.9)
    tbl.columns[1].width = Inches(2.9)

    # Header row — amethyst background, white bold label
    for cell, label in zip(tbl.rows[0].cells, ['ENGLISH', 'BAHASA INDONESIA']):
        _set_cell_bg(cell, AMETHYST_HX)
        p = cell.paragraphs[0]
        p.clear()
        _para_spacing(p, before=80, after=80)
        r = p.add_run(label)
        r.font.name      = FONT_UI
        r.font.size      = SZ_LABEL
        r.font.bold      = True
        r.font.color.rgb = WHITE

    # One data row per bilingual block (one top-level clause)
    for block in blocks:
        data_row = tbl.add_row().cells
        en_cell, id_cell = data_row[0], data_row[1]

        # English side
        _render_lines_to_cell(en_cell, block['en'])

        # Indonesian side — engine-produced text or sworn-translator placeholder
        id_text = (block.get('id') or '').strip()
        if id_text:
            _render_lines_to_cell(id_cell, id_text)
        else:
            p = id_cell.paragraphs[0]
            p.clear()
            _para_spacing(p, before=80, after=80)
            r = p.add_run(_SWORN_PLACEHOLDER)
            r.font.name      = FONT_BODY
            r.font.size      = SZ_BODY
            r.font.italic    = True
            r.font.color.rgb = SECONDARY


def _build_clean(text: str, company_name: str, doc_type: str,
                 corridor: str, date_str: str) -> Document:
    doc_type_label = DOC_TYPE_LABELS.get(doc_type.lower(),
                     doc_type.replace('_', ' ').title())
    corridor_label = CORRIDOR_LABELS.get(corridor.upper(), corridor)
    gov_law        = CORRIDOR_GOV_LAW.get(corridor.upper(), "—")

    doc, sect = _new_doc()
    _setup_header_footer(sect, doc_type_label, corridor_label)
    _add_cover_block(doc, doc_type_label, corridor_label, gov_law, "Clean Version", date_str)
    _add_deal_profile_table(doc, corridor, doc_type, company_name, date_str)

    if corridor.upper() in _BILINGUAL_CORRIDORS:
        # ── SG-ID and any future dual-language corridor ───────────────────────
        _build_clean_bilingual(doc, text)
    else:
        # ── Standard single-column English layout (CN-SG, SG-MY, CN-MY …) ───
        flag_type = None

        for line in text.splitlines():
            s = line.strip()
            if not s:
                flag_type = None
                continue
            if RE_SEPARATOR.match(s):
                continue
            if RE_BOX_END.match(s):
                flag_type = None
                continue

            # ── Markdown structural elements ──────────────────────────────────────
            m_h = RE_MD_HEADING.match(s)
            if m_h:
                flag_type = None
                _write_md_heading(doc, len(m_h.group(1)), m_h.group(2).strip())
                continue
            if RE_MD_HR.match(s):
                _horizontal_rule(doc)
                continue
            m_bq = RE_MD_BLOCK.match(s)
            if m_bq:
                flag_type = None
                content = m_bq.group(1).strip()
                _add_flag_para(doc, content or s, _detect_flag(content) or 'note')
                continue
            # ─────────────────────────────────────────────────────────────────────

            # Box-drawing blocks
            if RE_BOX_START.match(s):
                flag_type = _detect_flag(s) or 'lawyer'
                _add_flag_para(doc, s, flag_type, is_label=True)
                continue
            if RE_BOX_MID.match(s) and flag_type:
                _add_flag_para(doc, s, flag_type)
                continue

            # Inline flags
            detected = _detect_flag(s)
            if detected:
                flag_type = detected
                _add_flag_para(doc, s, flag_type, is_label=True)
                continue
            if flag_type and not RE_TOP_CLAUSE.match(s):
                _add_flag_para(doc, s, flag_type)
                continue
            flag_type = None

            # Clause headings
            level, num_str, title_str = _clause_parts(s)
            if level:
                _add_clause_heading(doc, num_str, title_str, level)
                continue

            _para_with_placeholders(doc, s)

    _add_execution_block(doc)
    return doc


# ══════════════════════════════════════════════════════════════════════════════
# REDLINE document
# ══════════════════════════════════════════════════════════════════════════════

def _build_redline(text: str, doc_type: str, corridor: str,
                   date_str: str) -> Document:
    doc_type_label = DOC_TYPE_LABELS.get(doc_type.lower(),
                     doc_type.replace('_', ' ').title())
    corridor_label = CORRIDOR_LABELS.get(corridor.upper(), corridor)
    gov_law        = CORRIDOR_GOV_LAW.get(corridor.upper(), "—")

    doc, sect = _new_doc()
    _setup_header_footer(sect, doc_type_label, corridor_label)
    _add_cover_block(doc, doc_type_label, corridor_label, gov_law, "Redline Version", date_str)

    flag_type = None

    for line in text.splitlines():
        s = line.strip()
        if not s:
            flag_type = None
            continue
        if RE_SEPARATOR.match(s):
            continue
        if RE_BOX_END.match(s):
            flag_type = None
            continue

        # ── Markdown structural elements ──────────────────────────────────────
        m_h = RE_MD_HEADING.match(s)
        if m_h:
            flag_type = None
            _write_md_heading(doc, len(m_h.group(1)), m_h.group(2).strip())
            continue
        if RE_MD_HR.match(s):
            _horizontal_rule(doc)
            continue
        m_bq = RE_MD_BLOCK.match(s)
        if m_bq:
            flag_type = None
            content = m_bq.group(1).strip()
            _add_flag_para(doc, content or s, _detect_flag(content) or 'note')
            continue
        # ─────────────────────────────────────────────────────────────────────

        # ── Insertion — underlined green #059669 ──────────────────────────────
        if s.startswith('>> INSERTED:') or s.startswith('INSERTED:'):
            _, _, rest = s.partition(':')
            para = doc.add_paragraph()
            _para_spacing(para, after=80)
            rl = para.add_run("INSERTED: ")
            _fmt(rl, font=FONT_UI, bold=True, size=SZ_LABEL, color=INS_CLR)
            rb = para.add_run(rest.strip())
            _fmt(rb, underline=True, color=INS_CLR)
            continue

        # ── Deletion — strikethrough red #DC2626 ──────────────────────────────
        if (s.startswith('>> REMOVED:') or s.startswith('REMOVED:')
                or s.startswith('ORIGINAL:')):
            _, _, rest = s.partition(':')
            para = doc.add_paragraph()
            _para_spacing(para, after=80)
            rl = para.add_run("REMOVED: ")
            _fmt(rl, font=FONT_UI, bold=True, size=SZ_LABEL, color=DEL_CLR)
            rb = para.add_run(rest.strip())
            _fmt(rb, strike=True, color=DEL_CLR)
            continue

        # Flag boxes
        if RE_BOX_START.match(s):
            flag_type = _detect_flag(s) or 'lawyer'
            _add_flag_para(doc, s, flag_type, is_label=True)
            continue
        if RE_BOX_MID.match(s) and flag_type:
            _add_flag_para(doc, s, flag_type)
            continue
        detected = _detect_flag(s)
        if detected:
            flag_type = detected
            _add_flag_para(doc, s, flag_type, is_label=True)
            continue
        if flag_type and not RE_TOP_CLAUSE.match(s):
            _add_flag_para(doc, s, flag_type)
            continue
        flag_type = None

        # Clause headings
        level, num_str, title_str = _clause_parts(s)
        if level:
            _add_clause_heading(doc, num_str, title_str, level)
            continue

        para = doc.add_paragraph()
        _para_spacing(para, after=80)
        _add_md_runs(para, s)

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# COMMENTARY document
# ══════════════════════════════════════════════════════════════════════════════

def _action_color(text: str):
    u = text.upper()
    if 'LAWYER REVIEW' in u:           return LAWYER_CLR, True
    if 'ACTION REQUIRED' in u or 'USER INPUT' in u: return ACTION_CLR, False
    if 'VERIFY' in u:                  return NOTE_CLR, False
    if 'NONE' in u or 'FULL_AUTO' in u: return INS_CLR, False
    return BODY_CLR, False


_SEVERITY_STYLE = {
    'HIGH':   (HIGH_BG_HX,       DEL_CLR),
    'MEDIUM': (ACTION_BG_HX,     ACTION_CLR),
    'LOW':    (LOW_BG_HX,        SECONDARY),
}


def _render_summary_table(doc, rows: list):
    """Render the === SUMMARY TABLE === block as a colour-coded Word table."""
    # Section heading
    p_hdr = doc.add_paragraph()
    _para_spacing(p_hdr, before=200, after=80)
    _add_left_border(p_hdr, AMETHYST_HX, bg_hex=COVER_BG_HX, thickness='18')
    r_hdr = p_hdr.add_run("COMPLIANCE SUMMARY")
    r_hdr.font.name       = FONT_UI
    r_hdr.font.bold       = True
    r_hdr.font.small_caps = True
    r_hdr.font.size       = SZ_H1
    r_hdr.font.color.rgb  = AMETHYST

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(1.0)
    tbl.columns[1].width = Inches(4.2)
    tbl.columns[2].width = Inches(1.1)

    # Header row
    for cell, label in zip(tbl.rows[0].cells, ['Clause', 'Issue', 'Severity']):
        _set_cell_bg(cell, TABLE_HDR_HX)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(label)
        r.font.name  = FONT_UI
        r.font.size  = SZ_LABEL
        r.font.bold  = True
        r.font.color.rgb = AMETHYST

    for clause, issue, severity in rows:
        sev_key = severity.upper().strip()
        bg_hex, clr = _SEVERITY_STYLE.get(sev_key, (LOW_BG_HX, SECONDARY))
        data_row = tbl.add_row().cells
        for i, (cell, txt) in enumerate(zip(data_row, [clause, issue, severity])):
            _set_cell_bg(cell, bg_hex)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(txt)
            if i == 2:  # Severity column — UI font, bold, colour
                r.font.name  = FONT_UI
                r.font.size  = SZ_LABEL
                r.font.bold  = True
                r.font.color.rgb = clr
            else:
                r.font.name  = FONT_BODY
                r.font.size  = SZ_BODY
                r.font.color.rgb = BODY_CLR

    doc.add_paragraph()  # breathing room after table


def _render_pdpa_table(doc, rows: list):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    _set_cell_bg(hdr[0], TABLE_HDR_HX)
    _set_cell_bg(hdr[1], TABLE_HDR_HX)
    for cell, lbl in zip(hdr, ['Data Protection Obligation', 'Status']):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(lbl)
        r.font.name  = FONT_UI
        r.font.size  = SZ_LABEL
        r.font.bold  = True
        r.font.color.rgb = AMETHYST

    for row_text in rows:
        is_risk = row_text.startswith('PDPC enforcement risk')
        item, _, status = (row_text.rpartition(' — ')
                           if ' — ' in row_text else (row_text, '', ''))
        item   = item.lstrip('✅⚠️☑ ').strip()
        status = status.strip()
        u = status.upper()
        clr = (INS_CLR    if any(x in u for x in ('INSERTED', 'COMPLIANT', 'LOW'))   else
               ACTION_CLR if any(x in u for x in ('PARAMETERIZED', 'MEDIUM'))        else
               LAWYER_CLR if any(x in u for x in ('LAWYER', 'HIGH', 'REVIEW'))       else
               BODY_CLR)
        row = tbl.add_row().cells
        for cell, txt in zip(row, [item, status]):
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(txt)
            _fmt(r, bold=is_risk, color=clr)


def _build_commentary(text: str, company_name: str, doc_type: str,
                       corridor: str, date_str: str) -> Document:
    doc_type_label = DOC_TYPE_LABELS.get(doc_type.lower(),
                     doc_type.replace('_', ' ').title())
    corridor_label = CORRIDOR_LABELS.get(corridor.upper(), corridor)
    gov_law        = CORRIDOR_GOV_LAW.get(corridor.upper(), "—")

    doc, sect = _new_doc()
    _setup_header_footer(sect, doc_type_label, corridor_label)
    _add_cover_block(doc, doc_type_label, corridor_label, gov_law,
                     "Commentary & Legal Basis", date_str)
    _add_deal_profile_table(doc, corridor, doc_type, company_name, date_str)

    in_pdpa        = False
    pdpa_rows      = []
    flag_type      = None
    in_summary     = False
    summary_rows   = []

    for line in text.splitlines():
        s = line.strip()
        if not s:
            if not in_summary:
                flag_type = None
            continue

        # ── Summary table block ───────────────────────────────────────────────
        if RE_SUMMARY_TABLE_START.match(s):
            in_summary   = True
            summary_rows = []
            continue
        if RE_SUMMARY_TABLE_END.match(s):
            in_summary = False
            _render_summary_table(doc, summary_rows)
            continue
        if in_summary:
            m = RE_TABLE_ROW.match(s)
            if m:
                clause, issue, severity = m.group(1), m.group(2), m.group(3)
                if clause.lower().strip() != 'clause':   # skip header row
                    summary_rows.append((clause, issue, severity))
            continue
        # ─────────────────────────────────────────────────────────────────────

        # ── Markdown structural elements ──────────────────────────────────────
        m_h = RE_MD_HEADING.match(s)
        if m_h:
            flag_type = None
            _write_md_heading(doc, len(m_h.group(1)), m_h.group(2).strip())
            continue
        if RE_MD_HR.match(s):
            _horizontal_rule(doc)
            continue
        m_bq = RE_MD_BLOCK.match(s)
        if m_bq:
            flag_type = None
            content = m_bq.group(1).strip()
            _add_flag_para(doc, content or s, _detect_flag(content) or 'note')
            continue
        # ─────────────────────────────────────────────────────────────────────

        # PDPA checklist
        if 'PDPA COMPLIANCE CHECKLIST' in s or 'DATA PROTECTION CHECKLIST' in s:
            in_pdpa = True
            p = doc.add_paragraph()
            _para_spacing(p, before=240, after=80)
            _add_left_border(p, AMETHYST_HX, bg_hex=COVER_BG_HX, thickness='18')
            r = p.add_run(s)
            r.font.name       = FONT_UI
            r.font.bold       = True
            r.font.small_caps = True
            r.font.size       = SZ_H1
            r.font.color.rgb  = AMETHYST
            continue

        if in_pdpa:
            if RE_SEPARATOR.match(s):
                if pdpa_rows:
                    _render_pdpa_table(doc, pdpa_rows)
                    pdpa_rows = []
                in_pdpa = False
                continue
            if RE_PDPA_ITEM.match(s) or s.startswith('PDPC enforcement risk'):
                pdpa_rows.append(s)
            continue

        if RE_SEPARATOR.match(s):
            continue

        # Clause entry header — [CLAUSE X.X] ID in monospace + amethyst left border
        if RE_CLAUSE_HEADER.match(s):
            _horizontal_rule(doc)
            p = doc.add_paragraph()
            _para_spacing(p, before=160, after=60)
            _add_left_border(p, AMETHYST_HX, bg_hex=COVER_BG_HX, thickness='18')
            r_id = p.add_run(s)
            r_id.font.name  = FONT_MONO
            r_id.font.bold  = True
            r_id.font.size  = SZ_H1
            r_id.font.color.rgb = AMETHYST
            continue

        # Flag boxes
        if RE_BOX_END.match(s):
            flag_type = None
            continue
        if RE_BOX_START.match(s):
            flag_type = _detect_flag(s) or 'lawyer'
            _add_flag_para(doc, s, flag_type, is_label=True)
            continue
        if RE_BOX_MID.match(s) and flag_type:
            _add_flag_para(doc, s, flag_type)
            continue
        detected = _detect_flag(s)
        if detected:
            flag_type = detected
            _add_flag_para(doc, s, flag_type, is_label=True)
            continue
        if flag_type and not RE_CLAUSE_HEADER.match(s):
            _add_flag_para(doc, s, flag_type)
            continue
        flag_type = None

        # Four structured fields:
        #   Original:  → grey italic Times New Roman (original text)
        #   Change:    → dark Times New Roman
        #   Reason:    → blue-grey italic Times New Roman
        #   Action required: → colour-coded
        matched_lbl = next((l for l in FIELD_LABELS if s.startswith(l)), None)
        if matched_lbl:
            rest = s[len(matched_lbl):].strip()
            para = doc.add_paragraph()
            _para_spacing(para, before=0, after=40)
            _set_indent(para, left_twips=180)

            r_lbl = para.add_run(matched_lbl + '  ')
            r_lbl.font.name  = FONT_UI
            r_lbl.font.bold  = True
            r_lbl.font.size  = SZ_LABEL
            r_lbl.font.color.rgb = SECONDARY

            if matched_lbl == 'Original:':
                _add_md_runs(para, rest, italic=True, color=SECONDARY)
            elif matched_lbl == 'Reason:':
                _add_md_runs(para, rest, italic=True, color=BLUE_GREY)
            elif matched_lbl == 'Action required:':
                clr, bold = _action_color(rest)
                r_val = para.add_run(rest)
                _fmt(r_val, font=FONT_UI, bold=bold, size=SZ_LABEL, color=clr)
            else:
                _add_md_runs(para, rest)
            continue

        # Recommendations
        if s.startswith('→') or s.startswith('Recommendation:'):
            para = doc.add_paragraph()
            _para_spacing(para, before=0, after=40)
            _set_indent(para, left_twips=360)
            _add_md_runs(para, s, italic=True, color=SECONDARY)
            continue

        # Default body
        para = doc.add_paragraph()
        _para_spacing(para, after=80)
        _add_md_runs(para, s)

    if pdpa_rows:
        _render_pdpa_table(doc, pdpa_rows)

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# Public entry point — signature preserved from original
# ══════════════════════════════════════════════════════════════════════════════

def build_outputs(
    clean_text: str,
    redline_text: str,
    commentary_text: str,
    company_name: str,
    doc_type: str,
    output_dir: str,
    job_id: str = None,
    corridor: str = "CN_SG",
) -> dict:
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    prefix = f"{job_id}_" if job_id else ""

    date_str = datetime.utcnow().strftime("%d %B %Y")
    commentary_text = re.sub(
        r'(?i)Date of adaptation:[^\n]*',
        f'Date of adaptation: {date_str}',
        commentary_text,
    )

    paths = {}

    doc = _build_clean(clean_text, company_name, doc_type, corridor, date_str)
    paths['clean'] = str(out / f"{prefix}clean.docx")
    doc.save(paths['clean'])

    doc = _build_redline(redline_text, doc_type, corridor, date_str)
    paths['redline'] = str(out / f"{prefix}redline.docx")
    doc.save(paths['redline'])

    doc = _build_commentary(commentary_text, company_name, doc_type, corridor, date_str)
    paths['commentary'] = str(out / f"{prefix}commentary.docx")
    doc.save(paths['commentary'])

    return paths
